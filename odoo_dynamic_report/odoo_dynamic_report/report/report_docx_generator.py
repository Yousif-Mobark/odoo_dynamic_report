# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.

from odoo import models, api, _
from odoo.exceptions import UserError, ValidationError
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from io import BytesIO
import base64
import re
import logging

_logger = logging.getLogger(__name__)


class ReportDocxGenerator(models.AbstractModel):
    _name = 'report.docx.generator'
    _description = 'DOCX Report Generator'

    @api.model
    def generate_report(self, template, record_ids):
        """
        Generate DOCX report from template for given records
        
        Args:
            template: report.template record
            record_ids: list of record IDs to generate report for
            
        Returns:
            bytes: Generated DOCX file content
        """
        if not template.template_data:
            raise UserError(_("Template file is missing"))
        
        # Load the template
        template_bytes = base64.b64decode(template.template_data)
        doc = Document(BytesIO(template_bytes))
        
        # Get records
        model = self.env[template.model_name]
        records = model.browse(record_ids)
        
        if not records:
            raise UserError(_("No records found to generate report"))
        
        _logger.info(f"Generating report for {len(records)} record(s)")
        
        # Process each record
        if len(records) == 1:
            # Single record - fill the template
            self._fill_template(doc, records[0], template)
        else:
            # Multiple records - duplicate template for each
            self._fill_template_multiple(doc, records, template)
        
        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.read()

    def _fill_template(self, doc, record, template):
        """Fill template with single record data"""
        # Process paragraphs
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, record, template)
        
        # Process tables
        for table in doc.tables:
            self._process_table(table, record, template)
        
        # Process headers and footers
        for section in doc.sections:
            self._process_section(section, record, template)

    def _fill_template_multiple(self, doc, records, template):
        """Fill template with multiple records (page per record)"""
        from docx.oxml import OxmlElement
        
        # Process first record
        self._fill_template(doc, records[0], template)
        
        # For each additional record, add page break and duplicate content
        for record in records[1:]:
            # Add page break
            doc.add_page_break()
            
            # Process new record
            self._fill_template(doc, record, template)

    def _process_paragraph(self, paragraph, record, template):
        """Process a single paragraph and replace placeholders"""
        if not paragraph.text:
            return
        
        # Find all placeholders {{field_name}}
        text = paragraph.text
        placeholders = re.findall(r'\{\{([^}]+)\}\}', text)
        
        for placeholder in placeholders:
            # Get field value
            value = self._get_field_value(record, placeholder.strip(), template)
            
            # Replace placeholder
            text = text.replace(f'{{{{{placeholder}}}}}', str(value))
        
        # Update paragraph text if changed
        if text != paragraph.text:
            # Preserve formatting by replacing runs
            if paragraph.runs:
                paragraph.runs[0].text = text
                for run in paragraph.runs[1:]:
                    run.text = ''
            else:
                paragraph.text = text

    def _process_table(self, table, record, template):
        """Process table and handle loops for one2many fields"""
        for row in table.rows:
            # Check if this is a loop row {{#field_name}}
            row_text = ' '.join([cell.text for cell in row.cells])
            
            loop_match = re.search(r'\{\{#(\w+)\}\}', row_text)
            if loop_match:
                field_name = loop_match.group(1)
                self._process_table_loop(table, row, record, field_name, template)
            else:
                # Normal row - just replace placeholders
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, record, template)

    def _process_table_loop(self, table, template_row, record, field_name, template):
        """Process one2many field in table (duplicate rows)"""
        if not hasattr(record, field_name):
            return
        
        related_records = getattr(record, field_name)
        if not related_records:
            return
        
        # Get row index
        row_index = table.rows.index(template_row)
        
        # Process each related record
        for idx, related_record in enumerate(related_records):
            if idx > 0:
                # Duplicate row
                new_row = self._duplicate_table_row(table, row_index + idx)
            else:
                new_row = template_row
            
            # Fill row with related record data
            for cell in new_row.cells:
                for paragraph in cell.paragraphs:
                    # Remove loop markers
                    text = paragraph.text.replace(f'{{{{#{field_name}}}}}', '')
                    text = text.replace(f'{{{{/{field_name}}}}}', '')
                    paragraph.text = text
                    
                    # Process placeholders with related record
                    self._process_paragraph(paragraph, related_record, template)

    def _duplicate_table_row(self, table, row_index):
        """Duplicate a table row"""
        from copy import deepcopy
        
        row = table.rows[row_index]
        new_row = table.add_row()
        
        # Copy cell content
        for idx, cell in enumerate(row.cells):
            new_cell = new_row.cells[idx]
            new_cell.text = cell.text
            
            # Copy formatting
            for paragraph in cell.paragraphs:
                new_paragraph = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
        
        return new_row

    def _process_section(self, section, record, template):
        """Process headers and footers"""
        if section.header:
            for paragraph in section.header.paragraphs:
                self._process_paragraph(paragraph, record, template)
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                self._process_paragraph(paragraph, record, template)

    def _get_field_value(self, record, field_path, template):
        """Get field value from record using field path"""
        try:
            # Check if there's a custom formatter
            if '|' in field_path:
                field_path, formatter = field_path.split('|', 1)
                value = self._traverse_field_path(record, field_path.strip())
                return self._format_value(value, formatter.strip())
            else:
                value = self._traverse_field_path(record, field_path)
                return self._format_value(value, None)
        except Exception as e:
            _logger.warning(f"Error getting field value for {field_path}: {e}")
            return f"[Error: {field_path}]"

    def _traverse_field_path(self, record, field_path):
        """Traverse field path like 'partner_id.country_id.name'"""
        parts = field_path.split('.')
        value = record
        
        for part in parts:
            if not value:
                return ''
            
            if not hasattr(value, part):
                return ''
            
            value = getattr(value, part)
            
            # Handle recordsets - take first or display_name
            if isinstance(value, models.BaseModel):
                if len(value) > 1:
                    return ', '.join(value.mapped('display_name'))
                elif len(value) == 0:
                    return ''
        
        return value

    def _format_value(self, value, formatter):
        """Format value based on formatter string"""
        if value is False or value is None:
            return ''
        
        # Handle model recordsets
        if isinstance(value, models.BaseModel):
            return value.display_name if len(value) == 1 else ', '.join(value.mapped('display_name'))
        
        # Apply formatter
        if formatter:
            if formatter.startswith('date:'):
                from datetime import datetime
                date_format = formatter.replace('date:', '').strip("'\"")
                if isinstance(value, str):
                    try:
                        value = datetime.strptime(value, '%Y-%m-%d')
                    except:
                        pass
                if hasattr(value, 'strftime'):
                    return value.strftime(date_format)
            elif formatter.startswith('number:'):
                number_format = formatter.replace('number:', '').strip("'\"")
                try:
                    return format(float(value), number_format)
                except:
                    pass
            elif formatter == 'upper':
                return str(value).upper()
            elif formatter == 'lower':
                return str(value).lower()
            elif formatter == 'title':
                return str(value).title()
        
        return str(value)

    def _validate_template(self, template_data):
        """Validate that template_data is a valid DOCX file"""
        try:
            template_bytes = base64.b64decode(template_data)
            doc = Document(BytesIO(template_bytes))
            return True
        except Exception as e:
            raise ValidationError(_("Invalid DOCX file: %s") % str(e))

    def _extract_placeholders(self, template_data):
        """Extract all placeholders from template"""
        template_bytes = base64.b64decode(template_data)
        doc = Document(BytesIO(template_bytes))
        
        placeholders = set()
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
            placeholders.update([m.strip() for m in matches])
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
                        placeholders.update([m.strip() for m in matches])
        
        # Remove loop markers
        placeholders = {p for p in placeholders if not p.startswith('#') and not p.startswith('/')}
        
        return list(placeholders)
