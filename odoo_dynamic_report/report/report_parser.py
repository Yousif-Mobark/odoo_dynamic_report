# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.

from odoo import models, api, _
from odoo.exceptions import ValidationError
import re
import logging

_logger = logging.getLogger(__name__)


class ReportParser(models.AbstractModel):
    _name = 'report.parser'
    _description = 'Report Template Parser'

    @api.model
    def parse_template(self, template_content):
        """
        Parse template and extract all field references
        
        Args:
            template_content: Binary template content
            
        Returns:
            dict: Parsed template information
        """
        placeholders = self._extract_placeholders(template_content)
        structure = self._analyze_structure(template_content)
        
        return {
            'placeholders': placeholders,
            'structure': structure,
            'field_count': len(placeholders),
        }

    def _extract_placeholders(self, template_content):
        """Extract all {{field}} placeholders from template"""
        # This will be implemented by report.docx.generator
        return self.env['report.docx.generator']._extract_placeholders(template_content)

    def _analyze_structure(self, template_content):
        """Analyze template structure (tables, sections, etc.)"""
        from docx import Document
        from io import BytesIO
        import base64
        
        template_bytes = base64.b64decode(template_content)
        doc = Document(BytesIO(template_bytes))
        
        structure = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections),
            'tables': []
        }
        
        # Analyze tables
        for idx, table in enumerate(doc.tables):
            table_info = {
                'index': idx,
                'row_count': len(table.rows),
                'col_count': len(table.columns) if table.rows else 0,
                'has_loop': False,
            }
            
            # Check for loops in table
            for row in table.rows:
                row_text = ' '.join([cell.text for cell in row.cells])
                if re.search(r'\{\{#\w+\}\}', row_text):
                    table_info['has_loop'] = True
                    break
            
            structure['tables'].append(table_info)
        
        return structure

    @api.model
    def validate_field_path(self, model_name, field_path):
        """
        Validate that a field path exists on a model
        
        Args:
            model_name: Technical name of the model
            field_path: Field path like 'partner_id.name'
            
        Returns:
            dict: Validation result
        """
        try:
            model = self.env[model_name]
            parts = field_path.split('.')
            
            current_model = model
            field_chain = []
            
            for part in parts:
                if part not in current_model._fields:
                    return {
                        'valid': False,
                        'error': _("Field '%s' does not exist on model '%s'") % (part, current_model._name)
                    }
                
                field = current_model._fields[part]
                field_chain.append({
                    'name': part,
                    'type': field.type,
                    'string': field.string,
                    'model': current_model._name,
                })
                
                # Navigate to related model if relational field
                if field.type in ('many2one', 'one2many', 'many2many'):
                    if field.comodel_name:
                        current_model = self.env[field.comodel_name]
                    else:
                        return {
                            'valid': False,
                            'error': _("Relational field '%s' has no target model") % part
                        }
            
            return {
                'valid': True,
                'field_chain': field_chain,
                'final_type': field_chain[-1]['type'] if field_chain else None,
            }
            
        except KeyError:
            return {
                'valid': False,
                'error': _("Model '%s' does not exist") % model_name
            }
        except Exception as e:
            return {
                'valid': False,
                'error': str(e)
            }

    @api.model
    def get_available_fields(self, model_name, include_related=True, max_depth=2):
        """
        Get all available fields for a model
        
        Args:
            model_name: Technical name of the model
            include_related: Include related model fields
            max_depth: Maximum depth for related fields
            
        Returns:
            list: List of field information dictionaries
        """
        try:
            model = self.env[model_name]
            fields = []
            
            self._collect_fields(model, fields, [], max_depth if include_related else 0)
            
            return fields
            
        except KeyError:
            raise ValidationError(_("Model '%s' does not exist") % model_name)

    def _collect_fields(self, model, fields_list, path, remaining_depth):
        """Recursively collect fields from model"""
        for field_name, field in model._fields.items():
            # Skip internal fields
            if field_name.startswith('_') or field_name in ('id', 'create_uid', 'write_uid'):
                continue
            
            current_path = path + [field_name]
            field_path = '.'.join(current_path)
            
            field_info = {
                'name': field_name,
                'path': field_path,
                'string': field.string,
                'type': field.type,
                'model': model._name,
                'required': field.required,
                'readonly': field.readonly,
                'help': field.help or '',
                'depth': len(current_path) - 1,
            }
            
            # Add relation info
            if field.type in ('many2one', 'one2many', 'many2many'):
                field_info['relation'] = field.comodel_name
            
            fields_list.append(field_info)
            
            # Recursively collect related fields
            if remaining_depth > 0 and field.type in ('many2one',):
                try:
                    related_model = self.env[field.comodel_name]
                    self._collect_fields(
                        related_model,
                        fields_list,
                        current_path,
                        remaining_depth - 1
                    )
                except:
                    pass
