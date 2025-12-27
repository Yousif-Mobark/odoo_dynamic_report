# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.

from odoo.tests import common, tagged
from odoo.exceptions import ValidationError
import tempfile
import os
from docx import Document
from docx.shared import Inches


@tagged('post_install', '-at_install')
class TestReportParser(common.TransactionCase):
    """Test suite for report parser functionality"""

    def setUp(self):
        super(TestReportParser, self).setUp()
        self.parser = self.env['report.parser']
        
        # Create test partner model reference
        self.partner_model = self.env['ir.model'].search([
            ('model', '=', 'res.partner')
        ], limit=1)

    def _create_test_docx(self, placeholders):
        """Helper to create test DOCX file with placeholders"""
        doc = Document()
        
        # Add title
        doc.add_heading('Test Template', 0)
        
        # Add placeholders
        for placeholder in placeholders:
            doc.add_paragraph(placeholder)
        
        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name

    def test_simple_placeholder_extraction(self):
        """Test extraction of simple field placeholders"""
        test_file = self._create_test_docx([
            'Name: {{name}}',
            'Email: {{email}}',
            'Phone: {{phone}}'
        ])
        
        try:
            placeholders = self.parser.extract_placeholders(test_file)
            
            self.assertIn('name', placeholders)
            self.assertIn('email', placeholders)
            self.assertIn('phone', placeholders)
        finally:
            os.unlink(test_file)

    def test_nested_field_extraction(self):
        """Test extraction of nested field placeholders"""
        test_file = self._create_test_docx([
            'Partner: {{partner_id.name}}',
            'Country: {{country_id.name}}',
            'State: {{state_id.name}}'
        ])
        
        try:
            placeholders = self.parser.extract_placeholders(test_file)
            
            self.assertIn('partner_id.name', placeholders)
            self.assertIn('country_id.name', placeholders)
            self.assertIn('state_id.name', placeholders)
        finally:
            os.unlink(test_file)

    def test_formatted_field_extraction(self):
        """Test extraction of fields with formatters"""
        test_file = self._create_test_docx([
            'Date: {{create_date|date:"%Y-%m-%d"}}',
            'Amount: {{amount_total|number:",.2f"}}',
            'Name: {{name|upper}}'
        ])
        
        try:
            placeholders = self.parser.extract_placeholders(test_file)
            
            # Should extract base field names
            self.assertTrue(any('create_date' in p for p in placeholders))
            self.assertTrue(any('amount_total' in p for p in placeholders))
            self.assertTrue(any('name' in p for p in placeholders))
        finally:
            os.unlink(test_file)

    def test_table_loop_extraction(self):
        """Test extraction of table loop markers"""
        test_file = self._create_test_docx([
            '{{#order_line}}',
            'Product: {{product_id.name}}',
            'Quantity: {{product_uom_qty}}',
            '{{/order_line}}'
        ])
        
        try:
            placeholders = self.parser.extract_placeholders(test_file)
            
            # Should detect loop field
            self.assertTrue(any('order_line' in p for p in placeholders))
        finally:
            os.unlink(test_file)

    def test_field_path_validation(self):
        """Test field path validation against model"""
        # Valid field path
        is_valid = self.parser.validate_field_path(
            'name',
            self.partner_model.model
        )
        self.assertTrue(is_valid)
        
        # Valid nested field path
        is_valid = self.parser.validate_field_path(
            'country_id.name',
            self.partner_model.model
        )
        self.assertTrue(is_valid)
        
        # Invalid field path
        is_valid = self.parser.validate_field_path(
            'invalid_field_xyz',
            self.partner_model.model
        )
        self.assertFalse(is_valid)

    def test_available_fields_retrieval(self):
        """Test retrieval of available fields for a model"""
        fields = self.parser.get_available_fields(
            self.partner_model.model,
            max_depth=2
        )
        
        self.assertIsInstance(fields, list)
        self.assertTrue(len(fields) > 0)
        
        # Check that common partner fields are present
        field_names = [f['name'] for f in fields]
        self.assertIn('name', field_names)
        self.assertIn('email', field_names)
        self.assertIn('phone', field_names)

    def test_nested_field_depth_limit(self):
        """Test that nested field traversal respects depth limit"""
        # Get fields with depth limit
        fields_depth_1 = self.parser.get_available_fields(
            self.partner_model.model,
            max_depth=1
        )
        
        fields_depth_2 = self.parser.get_available_fields(
            self.partner_model.model,
            max_depth=2
        )
        
        # Depth 2 should have more fields (includes related model fields)
        self.assertGreater(len(fields_depth_2), len(fields_depth_1))

    def test_field_type_detection(self):
        """Test correct field type detection"""
        fields = self.parser.get_available_fields(
            self.partner_model.model,
            max_depth=1
        )
        
        # Find specific field and check its type
        name_field = next((f for f in fields if f['name'] == 'name'), None)
        self.assertIsNotNone(name_field)
        self.assertEqual(name_field['type'], 'char')
        
        email_field = next((f for f in fields if f['name'] == 'email'), None)
        self.assertIsNotNone(email_field)
        self.assertEqual(email_field['type'], 'char')

    def test_relational_field_handling(self):
        """Test handling of many2one, one2many, many2many fields"""
        fields = self.parser.get_available_fields(
            self.partner_model.model,
            max_depth=2
        )
        
        # Check for relational fields
        field_dict = {f['name']: f for f in fields}
        
        # country_id should be many2one
        if 'country_id' in field_dict:
            self.assertEqual(field_dict['country_id']['type'], 'many2one')

    def test_placeholder_with_spaces(self):
        """Test handling of placeholders with spaces"""
        test_file = self._create_test_docx([
            'Field: {{ name }}',  # Spaces around field name
            'Field2: {{email}}'   # No spaces
        ])
        
        try:
            placeholders = self.parser.extract_placeholders(test_file)
            
            # Should normalize and find both
            self.assertTrue(len(placeholders) >= 2)
        finally:
            os.unlink(test_file)

    def test_duplicate_placeholder_handling(self):
        """Test handling of duplicate placeholders"""
        test_file = self._create_test_docx([
            'Name: {{name}}',
            'Name again: {{name}}',
            'Name third time: {{name}}'
        ])
        
        try:
            placeholders = self.parser.extract_placeholders(test_file)
            
            # Should not duplicate
            name_count = sum(1 for p in placeholders if p == 'name')
            self.assertEqual(name_count, 1)
        finally:
            os.unlink(test_file)
